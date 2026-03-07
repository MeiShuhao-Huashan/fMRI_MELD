#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import math
import re
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple
from xml.etree import ElementTree as ET

import docx  # python-docx


def _norm_ws(s: str) -> str:
    s = s.replace("\u00a0", " ").replace("\n", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _norm_dash(s: str) -> str:
    return s.replace("–", "-").replace("—", "-").replace("−", "-")


def _norm_num_str(s: str) -> str:
    """
    Normalize minor manuscript formatting differences:
      - dash variants
      - whitespace
      - drop trailing .0 for simple decimals (10.0 -> 10)
    """
    s2 = _norm_dash(_norm_ws(s))
    # 10.0 -> 10; 33.0 -> 33 (but keep 0.3 etc)
    s2 = re.sub(r"(?<=\d)\.0(?!\d)", "", s2)
    return s2


def _extract_docx_tables(docx_path: Path) -> List[List[List[str]]]:
    d = docx.Document(str(docx_path))
    out: List[List[List[str]]] = []
    for t in d.tables:
        rows: List[List[str]] = []
        for row in t.rows:
            rows.append([_norm_ws(c.text) for c in row.cells])
        out.append(rows)
    return out


def _parse_markdown_table(md_path: Path) -> List[Dict[str, str]]:
    """
    Parse the first GitHub-flavoured markdown table found in a file into a list of dicts.
    """
    lines = md_path.read_text(encoding="utf-8").splitlines()
    # Find first header line like: | a | b |
    start = None
    for i, ln in enumerate(lines):
        if ln.strip().startswith("|") and ln.strip().endswith("|"):
            # Next line is the alignment row.
            if i + 1 < len(lines) and re.match(r"^\s*\|[-:| ]+\|\s*$", lines[i + 1]):
                start = i
                break
    if start is None:
        raise SystemExit(f"FAILED: no markdown table found in {md_path}")

    header = [h.strip() for h in lines[start].strip().strip("|").split("|")]
    out: List[Dict[str, str]] = []
    for ln in lines[start + 2 :]:
        if not (ln.strip().startswith("|") and ln.strip().endswith("|")):
            break
        parts = [p.strip() for p in ln.strip().strip("|").split("|")]
        if len(parts) != len(header):
            continue
        out.append({header[j]: parts[j] for j in range(len(header))})
    return out


def _parse_table2_maintext_md(md_path: Path) -> Dict[Tuple[str, str], Dict[str, str]]:
    """
    Return dict[(endpoint_key, model_key)] -> row dict.

    endpoint_key in {"boxdsc", "ppv", "pin"}
    model_key in {"t1", "fmri", "tracka"}
    """
    txt = md_path.read_text(encoding="utf-8")
    sections = {
        "boxdsc": "## A.",
        "ppv": "## B.",
        "pin": "## C.",
    }
    model_map = {
        "MELD (self-trained)": "t1",
        "rs-fMRI": "fmri",
        "TrackA (fusion)": "tracka",
    }
    out: Dict[Tuple[str, str], Dict[str, str]] = {}
    for ep_key, marker in sections.items():
        if marker not in txt:
            raise SystemExit(f"FAILED: missing section {marker!r} in {md_path}")
        # Slice from marker to next section or end.
        start = txt.index(marker)
        end = min([txt.find(m, start + 1) for m in ["## A.", "## B.", "## C."] if txt.find(m, start + 1) != -1] + [len(txt)])
        block = txt[start:end]
        # Write temp to parse first table in block.
        tmp = md_path.with_suffix(f".{ep_key}.tmp.md")
        try:
            tmp.write_text(block, encoding="utf-8")
            rows = _parse_markdown_table(tmp)
        finally:
            try:
                tmp.unlink()
            except Exception:
                pass
        for r in rows:
            mk = model_map.get(r.get("Model", "").strip(), None)
            if mk is None:
                continue
            out[(ep_key, mk)] = r
    return out


def _format_p_for_docx(p_str: str) -> str:
    s = p_str.strip()
    try:
        p = float(s)
    except Exception:
        # Already a formatted token (e.g., "NA")
        return s
    if p < 0.001:
        return "p<0.001"
    if p <= 0.05:
        return f"{p:.3f}"
    return f"{p:.2f}"


def _format_ci_for_docx(ci_str: str) -> str:
    """
    Convert "[1.9, 89.6]" or "[0.893, 1.12e+03]" -> "1.90-89.60" / "0.89-1120".
    """
    s = ci_str.strip()
    m = re.match(r"^\[(.*),(.*)\]$", s)
    if not m:
        return s
    lo_s = m.group(1).strip()
    hi_s = m.group(2).strip()

    def _fmt(x: str) -> str:
        try:
            v = float(x)
        except Exception:
            return x
        if v >= 1000:
            return str(int(round(v)))
        if v >= 10:
            return f"{v:.2f}"
        if v >= 1:
            return f"{v:.2f}"
        return f"{v:.3g}"

    return f"{_fmt(lo_s)}-{_fmt(hi_s)}"


def _condense_appeal_table1_row(cells: Sequence[str]) -> Tuple[str, str, str]:
    if len(cells) < 6:
        raise ValueError(f"Unexpected Table1 row cell count: {len(cells)}")
    # Appeal manuscript uses merged columns duplicated by python-docx.
    return (cells[0], cells[2], cells[4])

def _read_table1_expected_from_tsv(tsv_path: Path) -> List[Tuple[str, str, str]]:
    """
    Read expected Table 1 rows from the reproducible TSV output.

    Expected columns:
      - Characteristic
      - ILAE 1–2
      - ILAE 3–6
    """
    rows: List[Tuple[str, str, str]] = []
    with tsv_path.open(newline="") as f:
        r = csv.DictReader(f, delimiter="\t")
        fields = set(r.fieldnames or [])
        need = {"Characteristic", "ILAE 1–2", "ILAE 3–6"}
        if not need.issubset(fields):
            raise SystemExit(f"FAILED: Table1 TSV missing required columns {sorted(need)} in {tsv_path} (got {sorted(fields)})")
        for row in r:
            rows.append(
                (
                    _norm_ws(str(row.get("Characteristic", ""))),
                    _norm_ws(str(row.get("ILAE 1–2", ""))),
                    _norm_ws(str(row.get("ILAE 3–6", ""))),
                )
            )
    if not rows:
        raise SystemExit(f"FAILED: Table1 TSV has no data rows: {tsv_path}")
    return rows


def _read_table1_expected_from_revision_docx(docx_path: Path) -> List[Tuple[str, str, str]]:
    d = docx.Document(str(docx_path))
    if not d.tables:
        raise SystemExit(f"FAILED: no tables in expected Table1 docx: {docx_path}")
    t = d.tables[0]
    out: List[Tuple[str, str, str]] = []
    for row in t.rows:
        cells = [_norm_ws(c.text) for c in row.cells]
        if len(cells) != 3:
            raise SystemExit(f"FAILED: expected Table1 template to have 3 cols, got {len(cells)} in row {cells}")
        out.append((cells[0], cells[1], cells[2]))
    return out


def _table1_cell_matches(appeal_cell: str, expected_cell: str) -> bool:
    a = _norm_num_str(appeal_cell)
    e_full = _norm_num_str(expected_cell)
    if a == e_full:
        return True
    # Allow appeal to keep only the first segment of a semicolon-delimited expected cell.
    if ";" in e_full:
        first = _norm_num_str(e_full.split(";", 1)[0].strip())
        if a == first:
            return True
    # Allow minor rounding differences (e.g., 2.48 vs 2.478) by comparing extracted numeric tokens.
    num_re = r"[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?"
    a_nums = [float(x) for x in re.findall(num_re, a)]
    e_nums = [float(x) for x in re.findall(num_re, e_full)]
    if a_nums and (len(a_nums) == len(e_nums)):
        if all(abs(aa - ee) <= 0.01 for aa, ee in zip(a_nums, e_nums)):
            return True
    return False


def _xlsx_shared_strings(z: zipfile.ZipFile) -> List[str]:
    xml = z.read("xl/sharedStrings.xml")
    root = ET.fromstring(xml)
    ns = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    out: List[str] = []
    for si in root.findall("m:si", ns):
        parts = [t.text or "" for t in si.findall(".//m:t", ns)]
        out.append("".join(parts))
    return out


def _xlsx_sheet_paths(z: zipfile.ZipFile) -> Dict[str, str]:
    """
    Map sheet name -> worksheet xml path inside the zip.
    """
    ns_wb = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main", "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
    wb = ET.fromstring(z.read("xl/workbook.xml"))
    sheets = wb.find("m:sheets", ns_wb)
    if sheets is None:
        raise ValueError("workbook.xml missing <sheets>")
    rid_by_name: Dict[str, str] = {}
    for sh in sheets.findall("m:sheet", ns_wb):
        name = str(sh.attrib.get("name", ""))
        rid = str(sh.attrib.get(f"{{{ns_wb['r']}}}id", ""))
        if name and rid:
            rid_by_name[name] = rid

    rels = ET.fromstring(z.read("xl/_rels/workbook.xml.rels"))
    ns_rel = {"m": "http://schemas.openxmlformats.org/package/2006/relationships"}
    target_by_rid: Dict[str, str] = {}
    for rel in rels.findall("m:Relationship", ns_rel):
        rid = str(rel.attrib.get("Id", ""))
        tgt = str(rel.attrib.get("Target", ""))
        if rid and tgt:
            target_by_rid[rid] = tgt

    out: Dict[str, str] = {}
    for name, rid in rid_by_name.items():
        tgt = target_by_rid.get(rid, "")
        if tgt:
            out[name] = f"xl/{tgt}".replace("\\\\", "/")
    return out


def _xlsx_parse_dimension(ref: str) -> Tuple[int, int]:
    m = re.match(r"^[A-Z]+[0-9]+:([A-Z]+)([0-9]+)$", ref.strip())
    if not m:
        # Fallback: treat as a single cell like A1
        m2 = re.match(r"^[A-Z]+([0-9]+)$", ref.strip())
        if not m2:
            return (0, 0)
        return (0, int(m2.group(1)))
    col = m.group(1)
    row = int(m.group(2))
    n = 0
    for ch in col:
        n = n * 26 + (ord(ch) - ord("A") + 1)
    return (n, row)


def _xlsx_col_to_idx(col: str) -> int:
    n = 0
    for ch in col:
        n = n * 26 + (ord(ch.upper()) - ord("A") + 1)
    return n


def _xlsx_ref_to_rc(ref: str) -> Tuple[int, int]:
    col = "".join(ch for ch in ref if ch.isalpha())
    row = "".join(ch for ch in ref if ch.isdigit())
    return (int(row), _xlsx_col_to_idx(col))


def _xlsx_read_sheet_matrix(xlsx: Path, sheet_name: str) -> List[List[object]]:
    with zipfile.ZipFile(xlsx, "r") as z:
        shared = _xlsx_shared_strings(z)
        sheet_paths = _xlsx_sheet_paths(z)
        sheet_path = sheet_paths.get(sheet_name, "")
        if not sheet_path:
            raise SystemExit(f"FAILED: sheet {sheet_name!r} not found in {xlsx}")
        root = ET.fromstring(z.read(sheet_path))

    ns = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    dim = root.find("m:dimension", ns)
    max_col = 0
    max_row = 0
    if dim is not None and "ref" in dim.attrib:
        max_col, max_row = _xlsx_parse_dimension(dim.attrib["ref"])

    cells: Dict[Tuple[int, int], object] = {}
    sheet_data = root.find("m:sheetData", ns)
    if sheet_data is None:
        return []
    for row in sheet_data.findall("m:row", ns):
        for c in row.findall("m:c", ns):
            ref = str(c.attrib.get("r", ""))
            if not ref:
                continue
            r, cc = _xlsx_ref_to_rc(ref)
            max_row = max(max_row, r)
            max_col = max(max_col, cc)
            t = c.attrib.get("t", "")
            if t == "s":
                v = c.find("m:v", ns)
                if v is None or v.text is None:
                    val: object = ""
                else:
                    idx = int(v.text)
                    val = shared[idx] if 0 <= idx < len(shared) else ""
                cells[(r, cc)] = val
                continue
            if t == "inlineStr":
                texts = [t_el.text or "" for t_el in c.findall(".//m:t", ns)]
                cells[(r, cc)] = "".join(texts)
                continue
            v = c.find("m:v", ns)
            if v is None or v.text is None:
                cells[(r, cc)] = ""
                continue
            s = v.text.strip()
            try:
                cells[(r, cc)] = float(s)
            except Exception:
                cells[(r, cc)] = s

    out: List[List[object]] = []
    for r in range(1, max_row + 1):
        row_vals: List[object] = []
        for c in range(1, max_col + 1):
            row_vals.append(cells.get((r, c), ""))
        out.append(row_vals)
    return out


def _read_tsv(path: Path) -> List[Dict[str, str]]:
    with path.open(newline="") as f:
        r = csv.DictReader(f, delimiter="\t")
        return [{k: str(v) for k, v in row.items()} for row in r]


def _maybe_float(s: str) -> Optional[float]:
    ss = str(s).strip()
    if ss == "" or ss.lower() in {"na", "nan"}:
        return None
    try:
        return float(ss)
    except Exception:
        return None


def _assert_close(label: str, got: object, exp: str, *, tol: float = 1e-6) -> None:
    if isinstance(got, float):
        ev = _maybe_float(exp)
        if ev is None:
            raise SystemExit(f"FAILED: {label}: expected non-numeric {exp!r}, got numeric {got}")
        if not (math.isfinite(got) and math.isfinite(ev) and abs(got - ev) <= tol):
            raise SystemExit(f"FAILED: {label}: got {got} expected {ev} (from {exp!r})")
        return
    if _norm_ws(str(got)) != _norm_ws(str(exp)):
        raise SystemExit(f"FAILED: {label}: got {got!r} expected {exp!r}")


def _check_table2_docx(*, docx_tables: List[List[List[str]]], table2_md: Path) -> None:
    # Appeal manuscript: table index 1 is Table2.
    t = docx_tables[1]
    # expected from md (Holm-corrected p already present)
    exp = _parse_table2_maintext_md(table2_md)

    # Map appeal endpoint labels to our endpoint keys.
    endpoint_key = {_norm_ws(k): v for k, v in {
        "PPV-in-mask ≥ 0.5a": "ppv",
        "boxDSC >0.22b": "boxdsc",
        "Pinpointing (COM in RC)c": "pin",
    }.items()}
    model_key = {_norm_ws(k): v for k, v in {
        "T1-only": "t1",
        "rs-fMRI-only": "fmri",
        "Late fusion": "tracka",
    }.items()}

    # Data rows are 1..9
    for i in range(1, len(t)):
        row = t[i]
        if len(row) != 6:
            raise SystemExit(f"FAILED: unexpected Table2 row width {len(row)} at row {i}: {row}")
        ep = endpoint_key.get(_norm_ws(row[0]), None)
        mk = model_key.get(_norm_ws(row[1]), None)
        if ep is None or mk is None:
            raise SystemExit(f"FAILED: cannot map Table2 row {i}: {row}")
        r = exp[(ep, mk)]

        exp_il12 = r["Resection=Yes in ILAE12 n/N (%)"]
        exp_il36 = r["Resection=Yes in ILAE3456 n/N (%)"]
        exp_or = r["Adjusted aOR"]
        exp_p = _format_p_for_docx(r["p"])

        if _norm_ws(row[2]) != _norm_ws(exp_il12):
            raise SystemExit(f"FAILED: Table2 row {i} ILAE12 got {row[2]!r} expected {exp_il12!r}")
        if _norm_ws(row[3]) != _norm_ws(exp_il36):
            raise SystemExit(f"FAILED: Table2 row {i} ILAE36 got {row[3]!r} expected {exp_il36!r}")
        # OR(CI): compare numerically to avoid cosmetic formatting differences (e.g., 8.4 vs 8.40),
        # and to avoid mis-parsing the CI hyphen as a negative sign.
        num_re = r"[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?"
        m = re.match(
            rf"^\s*(?P<or>{num_re})\s*\(\s*(?P<lo>{num_re})\s*-\s*(?P<hi>{num_re})\s*\)\s*$",
            row[4],
        )
        if m is None:
            raise SystemExit(f"FAILED: Table2 row {i} OR(CI) parse failed: got_cell={row[4]!r}")
        got_or = float(m.group("or"))
        got_lo = float(m.group("lo"))
        got_hi = float(m.group("hi"))

        ci_nums = [float(x) for x in re.findall(num_re, r["95% CI"])]
        if len(ci_nums) < 2:
            raise SystemExit(f"FAILED: Table2 row {i} expected CI parse failed: {r['95% CI']!r}")
        exp_or_f = float(exp_or)
        exp_lo = float(ci_nums[0])
        exp_hi = float(ci_nums[1])

        for j, (g, e) in enumerate([(got_or, exp_or_f), (got_lo, exp_lo), (got_hi, exp_hi)]):
            if not math.isclose(g, e, rel_tol=1e-3, abs_tol=0.01):
                raise SystemExit(f"FAILED: Table2 row {i} OR(CI) mismatch at pos {j}: got {g} expected {e} (got_cell={row[4]!r})")
        if _norm_ws(row[5]) != _norm_ws(exp_p):
            raise SystemExit(f"FAILED: Table2 row {i} p got {row[5]!r} expected {exp_p!r}")


def _check_table3_docx(*, docx_tables: List[List[List[str]]], table3_md: Path) -> None:
    # Appeal manuscript: table index 2 is Table3.
    t = docx_tables[2]
    md_rows = _parse_markdown_table(table3_md)

    # Build expected by normalized key from MD "Metric" (and sometimes "Scale").
    exp_by_key: Dict[str, Dict[str, str]] = {}
    for r in md_rows:
        metric = _norm_ws(r["Metric"]).lower()
        scale = _norm_ws(r["Scale"]).lower()
        key = re.sub(r"[^a-z0-9]+", "_", f"{scale}_{metric}").strip("_")
        exp_by_key[key] = r

    def docx_row_key(level: str, metric: str) -> str:
        lvl = _norm_ws(level).lower()
        met = _norm_ws(metric).lower()
        if lvl == "vertex":
            return "vertex_level_dice_similarity_coefficient_dsc_mean_95_ci"
        if "cluster" in lvl:
            if met.startswith("number"):
                return "cluster_level_boxdsc_0_22_dice_0_01_predicted_clusters_per_subject_mean_95_ci"
            if met.startswith("fp"):
                return "cluster_level_boxdsc_0_22_dice_0_01_false_positive_clusters_per_subject_mean_95_ci"
            if met.startswith("precision"):
                return "cluster_level_boxdsc_0_22_dice_0_01_precision_tp_clusters_all_clusters_95_ci"
            if met.startswith("f1"):
                return "cluster_level_boxdsc_0_22_dice_0_01_f1_score_mean_95_ci"
        if lvl == "subject" or lvl == "":
            if "boxdsc" in met and "detection" in met:
                return "subject_level_primary_detection_rate_det_boxdsc_0_22_dice_0_01_n_n_95_ci"
            if "ppv" in met and "detection" in met:
                return "subject_level_co_primary_detection_rate_det_ppv_in_mask_0_5_n_n_95_ci"
            if met.startswith("pinpointing"):
                return "subject_level_pinpointing_rate_any_cluster_com_within_lesion_n_n_95_ci"
            if "dist" in met and "detection" in met:
                return "subject_level_detection_rate_det_distance_20mm_n_n_95_ci"
            if met.startswith("recall"):
                return "subject_level_recall_0_15_lesion_dice_union_0_15_n_n_95_ci"
        return re.sub(r"[^a-z0-9]+", "_", f"{lvl}_{met}").strip("_")

    # Validate header row shape.
    if len(t) < 2 or len(t[0]) < 8:
        raise SystemExit("FAILED: Table3 structure unexpected in docx.")

    current_level = ""
    for i in range(1, len(t)):
        row = t[i]
        if len(row) != 8:
            raise SystemExit(f"FAILED: unexpected Table3 row width {len(row)} at row {i}: {row}")
        if _norm_ws(row[0]):
            current_level = row[0]
        key = docx_row_key(current_level, row[1])
        exp = exp_by_key.get(key, None)
        if exp is None:
            raise SystemExit(f"FAILED: cannot map Table3 row {i} to expected key: {key} row={row}")

        # Column mapping from MD to docx:
        # - docx[2] T1-only
        # - docx[3] rs-fMRI-only
        # - docx[4] Late fusion -> Track A (T1+fMRI)
        # - docx[5] Δ
        # - docx[6] p
        # - docx[7] p_adj
        if _norm_ws(row[2]) != _norm_ws(exp["T1-only"]):
            raise SystemExit(f"FAILED: Table3 row {i} T1 got {row[2]!r} expected {exp['T1-only']!r}")
        if _norm_ws(row[3]) != _norm_ws(exp["fMRI-only (fMRIhemi B2a)"]):
            raise SystemExit(f"FAILED: Table3 row {i} fMRI got {row[3]!r} expected {exp['fMRI-only (fMRIhemi B2a)']!r}")
        if _norm_ws(row[4]) != _norm_ws(exp["Track A (T1+fMRI)"]):
            raise SystemExit(f"FAILED: Table3 row {i} fusion got {row[4]!r} expected {exp['Track A (T1+fMRI)']!r}")
        if _norm_ws(row[5]) != _norm_ws(exp["Δ (Track A (T1+fMRI) − T1-only)"]):
            raise SystemExit(f"FAILED: Table3 row {i} delta got {row[5]!r} expected {exp['Δ (Track A (T1+fMRI) − T1-only)']!r}")
        # Manuscript Table3 prints only the p-value, while the markdown includes test details in parentheses.
        exp_p = _norm_ws(exp["p (Track A (T1+fMRI) vs T1-only)"]).split(" ", 1)[0]
        if _norm_ws(row[6]) != exp_p:
            raise SystemExit(f"FAILED: Table3 row {i} p got {row[6]!r} expected {exp_p!r}")

        # p_adj: compare numerically to allow '1' vs '1.00'.
        try:
            got_adj = float(_norm_ws(row[7]))
            exp_adj = float(_norm_ws(str(exp["p_adj (multiplicity)"])))
        except Exception:
            raise SystemExit(f"FAILED: Table3 row {i} p_adj parse failed: got {row[7]!r} expected {exp['p_adj (multiplicity)']!r}")
        if not math.isclose(got_adj, exp_adj, rel_tol=0.0, abs_tol=1e-9):
            raise SystemExit(f"FAILED: Table3 row {i} p_adj got {row[7]!r} expected {exp['p_adj (multiplicity)']!r}")


def _check_table1_docx(*, docx_tables: List[List[List[str]]], exp_rows: List[Tuple[str, str, str]]) -> None:
    # Appeal manuscript: table index 0 is Table1.
    t = docx_tables[0]
    # Drop rows that are omitted in the appeal manuscript (all-zero categories).
    exp_rows = [r for r in exp_rows if r[0] not in {"Easy", "Polymicrogyria"}]

    got_rows: List[Tuple[str, str, str]] = []
    for row in t:
        got_rows.append(_condense_appeal_table1_row(row))

    # Skip header row ("Characteristic").
    if exp_rows and _norm_ws(exp_rows[0][0]).lower() == "characteristic":
        exp_rows = exp_rows[1:]
    if got_rows and _norm_ws(got_rows[0][0]).lower() == "characteristic":
        got_rows = got_rows[1:]

    if len(got_rows) != len(exp_rows):
        raise SystemExit(f"FAILED: Table1 row count mismatch: got {len(got_rows)} expected {len(exp_rows)}")

    for i, (g, e) in enumerate(zip(got_rows, exp_rows)):
        g_lab, g12, g36 = g
        e_lab, e12, e36 = e

        # Normalize labels (ignore footnote markers and asterisks).
        def _norm_label(x: str) -> str:
            x2 = _norm_ws(x).replace("*", "")
            x2 = re.sub(r", n \(%\)[A-Za-z]$", ", n (%)", x2)
            x2 = re.sub(r"\)[a-zA-Z]$", ")", x2)
            x2 = re.sub(r"Encephalomalacia[a-zA-Z]$", "Encephalomalacia", x2)
            x2 = re.sub(r"rs-fMRI QC summary[a-zA-Z]$", "rs-fMRI QC summary", x2)
            return x2

        if _norm_label(g_lab) != _norm_label(e_lab):
            raise SystemExit(f"FAILED: Table1 label mismatch at row {i}: got {g_lab!r} expected {e_lab!r}")
        if not _table1_cell_matches(g12, e12):
            raise SystemExit(f"FAILED: Table1 ILAE12 mismatch at row {i} ({g_lab}): got {g12!r} expected {e12!r}")
        if not _table1_cell_matches(g36, e36):
            raise SystemExit(f"FAILED: Table1 ILAE36 mismatch at row {i} ({g_lab}): got {g36!r} expected {e36!r}")


def _check_supp_xlsx(*, supp_xlsx: Path, s2_tsv: Path, s3_tsv: Path, s4_tsv: Path) -> None:
    # ---- S2 ----
    mat = _xlsx_read_sheet_matrix(supp_xlsx, "Supplementary_TableS2")
    header = [str(x) for x in mat[0]]
    exp_rows = _read_tsv(s2_tsv)
    exp_by_test = {r["Index test"]: r for r in exp_rows}
    idx_col = header.index("Index test")
    test_name_map = {
        "T1 (MELD self-trained)": "MELD (self-trained)",
        "Late fusion": "TrackA (fusion)",
    }
    for r_idx, row in enumerate(mat[1:], start=2):
        test_raw = str(row[idx_col])
        test = test_name_map.get(test_raw, test_raw)
        if test not in exp_by_test:
            raise SystemExit(
                f"FAILED: Supplementary_TableS2 unexpected Index test {test_raw!r} (mapped={test!r}) at xlsx row {r_idx}"
            )
        exp = exp_by_test[test]
        for c_idx, col_name in enumerate(header):
            if not col_name:
                continue
            if col_name == "Index test":
                continue
            if col_name not in exp:
                continue
            _assert_close(f"S2 {test} {col_name}", row[c_idx], exp[col_name])

    # ---- S3 ----
    mat3 = _xlsx_read_sheet_matrix(supp_xlsx, "Supplementary_TableS3")
    exp3 = s3_tsv.read_text(encoding="utf-8").splitlines()

    # Build expected rows as a list of list[str] using TSV lines (keep the two header blocks).
    exp3_rows: List[List[str]] = []
    for ln in exp3:
        if not ln.strip():
            continue
        exp3_rows.append([x.strip() for x in ln.split("\t")])

    # Convert xlsx matrix rows to comparable strings (numbers -> compact string).
    got3_rows: List[List[str]] = []
    for row in mat3:
        out_row: List[str] = []
        for v in row:
            if isinstance(v, float) and math.isfinite(v):
                # Match TSV formatting: keep up to 4 sig figs, but avoid trailing .0.
                s = f"{v:.10g}"
                out_row.append(s)
            else:
                out_row.append(_norm_ws(str(v)))
        got3_rows.append(out_row)

    if len(got3_rows) != len(exp3_rows):
        raise SystemExit(f"FAILED: Supplementary_TableS3 row count mismatch: got {len(got3_rows)} expected {len(exp3_rows)}")
    for i, (g, e) in enumerate(zip(got3_rows, exp3_rows), start=1):
        if len(g) != len(e):
            raise SystemExit(f"FAILED: Supplementary_TableS3 col count mismatch at row {i}: got {len(g)} expected {len(e)}")
        for j, (gv, ev) in enumerate(zip(g, e), start=1):
            # Float-compare if both look numeric.
            gv_f = _maybe_float(gv)
            ev_f = _maybe_float(ev)
            if gv_f is not None and ev_f is not None:
                if not (math.isfinite(gv_f) and math.isfinite(ev_f) and abs(gv_f - ev_f) <= 1e-6):
                    raise SystemExit(f"FAILED: Supplementary_TableS3 mismatch at row {i} col {j}: got {gv} expected {ev}")
            else:
                if _norm_ws(gv) != _norm_ws(ev):
                    raise SystemExit(f"FAILED: Supplementary_TableS3 mismatch at row {i} col {j}: got {gv!r} expected {ev!r}")

    # ---- S4 ----
    mat4 = _xlsx_read_sheet_matrix(supp_xlsx, "Supplementary_TableS4")
    exp4_rows = _read_tsv(s4_tsv)
    exp_by_id = {r["Experiment"]: r for r in exp4_rows}

    # Appeal xlsx S4 uses a fixed 16-col layout; compare by columns with our TSV mapping.
    # Column letters: A..P -> indices 0..15
    for r_i, row in enumerate(mat4[1:], start=2):
        exp_id = str(row[0]).strip()
        if exp_id not in exp_by_id:
            raise SystemExit(f"FAILED: Supplementary_TableS4 unexpected Experiment {exp_id!r} at xlsx row {r_i}")
        exp = exp_by_id[exp_id]

        # Numeric overall metrics
        mapping_numeric = {
            3: "Vertex DSC (mean)",
            4: "Det(boxDSC>0.22) (overall)",
            5: "Pinpointing (overall)",
            6: "FP clusters/pt (overall)",
            7: "Cluster precision (overall)",
            8: "Cluster F1 (overall)",
        }
        for idx, key in mapping_numeric.items():
            _assert_close(f"S4 {exp_id} {key}", row[idx], exp[key])

        # Rescue strings in J/K/L (xlsx) map to the updated headers in our TSV.
        rescue_det_key = "Rescue Det(boxDSC>0.22) among T1 Det(boxDSC) fails (n=17)"
        rescue_pin_key = "Rescue Pinpointing among T1 pinpointing fails within T1 Det(boxDSC) fails (n=14)"
        rescue_ppv_key = "Rescue Det(PPV>=0.5) among T1 Det(PPV) fails within T1 Det(boxDSC) fails (n=15)"
        _assert_close(f"S4 {exp_id} rescue_det", row[9], exp[rescue_det_key])
        _assert_close(f"S4 {exp_id} rescue_pin", row[10], exp[rescue_pin_key])
        _assert_close(f"S4 {exp_id} rescue_ppv", row[11], exp[rescue_ppv_key])

        # Ranks + mean rank in M/N/O/P
        _assert_close(f"S4 {exp_id} rank_det", row[12], exp["Rescue Det(boxDSC) rank (1=best)"])
        _assert_close(f"S4 {exp_id} rank_pin", row[13], exp["Rescue Pinpointing rank (1=best)"])
        _assert_close(f"S4 {exp_id} rank_ppv", row[14], exp["Rescue Det(PPV>=0.5) rank (1=best)"])
        _assert_close(f"S4 {exp_id} mean_rank", row[15], exp["Mean rescue rank"])


def main() -> int:
    ap = argparse.ArgumentParser(description="Check appeal-final manuscript + Supplementary_Tables.xlsx match the reproducible repo outputs.")
    ap.add_argument("--docx", required=True, type=Path, help="Path to the appeal-final manuscript .docx (outside the public release).")
    ap.add_argument("--supp_xlsx", required=True, type=Path, help="Path to appeal Supplementary_Tables.xlsx")

    # Expected outputs (within release repo)
    ap.add_argument(
        "--table1_expected_tsv",
        default=Path("paper/revision/table1/table1_outcome58_baseline.tsv"),
        type=Path,
        help="Expected Table1 output TSV from this reproducible repo.",
    )
    ap.add_argument(
        "--table1_revision_docx",
        default=None,
        type=Path,
        help="Optional legacy: a formatted Table1 .docx used as an expected template (not required for public release).",
    )
    ap.add_argument(
        "--table2_maintext_md",
        default=Path("paper/revision/table2/final/prognosis_table_maintext_selftrained_unconstrained_meld_tracka_paper_plus20.md"),
        type=Path,
    )
    ap.add_argument(
        "--table3_md",
        default=Path("paper/revision/table3/final/table3_scope38_multiplicity_adjusted_epilepsia.md"),
        type=Path,
    )
    ap.add_argument(
        "--supp_s2_tsv",
        default=Path("paper/revision/supplement/table3_ppv_continuous_selftrained.tsv"),
        type=Path,
    )
    ap.add_argument(
        "--supp_s3_tsv",
        default=Path("paper/revision/supplement/native_meld/table2_native_meld_inject_union_tuned_thr0.10_diff0.15.tsv"),
        type=Path,
    )
    ap.add_argument(
        "--supp_s4_tsv",
        default=Path("paper/supplement_fmri_ablation/tableS_fmri_ablation_matrix_scope38.tsv"),
        type=Path,
    )
    args = ap.parse_args()

    docx_path = args.docx.expanduser().resolve()
    supp_xlsx = args.supp_xlsx.expanduser().resolve()
    if not docx_path.is_file():
        raise SystemExit(f"Missing docx: {docx_path}")
    if not supp_xlsx.is_file():
        raise SystemExit(f"Missing xlsx: {supp_xlsx}")

    root = Path(__file__).resolve().parents[2]
    table1_expected_tsv = (root / args.table1_expected_tsv).resolve()
    table1_revision_docx = (root / args.table1_revision_docx).resolve() if args.table1_revision_docx else None
    table2_md = (root / args.table2_maintext_md).resolve()
    table3_md = (root / args.table3_md).resolve()
    s2_tsv = (root / args.supp_s2_tsv).resolve()
    s3_tsv = (root / args.supp_s3_tsv).resolve()
    s4_tsv = (root / args.supp_s4_tsv).resolve()

    for p in [table1_expected_tsv, table2_md, table3_md, s2_tsv, s3_tsv, s4_tsv]:
        if not p.exists():
            raise SystemExit(f"Missing expected file under release repo: {p}")
    if table1_revision_docx is not None and not table1_revision_docx.exists():
        raise SystemExit(f"Missing expected file under release repo: {table1_revision_docx}")

    docx_tables = _extract_docx_tables(docx_path)
    if len(docx_tables) != 3:
        raise SystemExit(f"FAILED: expected 3 tables in appeal docx, got {len(docx_tables)}")

    # Prefer TSV as the expected Table1 source (public-release friendly). Fall back to docx if provided.
    exp_table1_rows = _read_table1_expected_from_tsv(table1_expected_tsv)
    if not exp_table1_rows and table1_revision_docx is not None:
        exp_table1_rows = _read_table1_expected_from_revision_docx(table1_revision_docx)

    _check_table1_docx(docx_tables=docx_tables, exp_rows=exp_table1_rows)
    _check_table2_docx(docx_tables=docx_tables, table2_md=table2_md)
    _check_table3_docx(docx_tables=docx_tables, table3_md=table3_md)
    _check_supp_xlsx(supp_xlsx=supp_xlsx, s2_tsv=s2_tsv, s3_tsv=s3_tsv, s4_tsv=s4_tsv)

    print("ok: appeal-final manuscript tables and Supplementary_Tables.xlsx match reproducible repo outputs.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
